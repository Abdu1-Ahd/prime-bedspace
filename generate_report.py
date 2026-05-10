import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()
    
    # 1. Title Page
    title = doc.add_heading('Prime BedSpace\nHospital Patient Triage & Bed Allocator', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\nCL2006 OS Lab · Spring 2026 · FAST-NUCES CFD')
    doc.add_paragraph('Group Members:\n- Abdul Ahad Zawiar (24F-0727)\n- Member 2 (24F-0514)\n')
    doc.add_page_break()

    # 2. Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        "1. Introduction", "2. System Design", "3. Phase 1: Shell Scripts",
        "4. Phase 2: Process Management & IPC", "5. Phase 3: Threads & Synchronization",
        "6. Phase 4: Memory Management", "7. Testing", "8. Valgrind Output",
        "9. Challenges & Lessons Learned", "10. Individual Contribution"
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_page_break()

    # 3. Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph("This project simulates a hospital triage and bed allocation system. It manages patient flow from arrival to discharge using core operating system concepts.")
    doc.add_paragraph("Objectives:")
    doc.add_paragraph("- Automate patient triage using shell scripts.", style='List Bullet')
    doc.add_paragraph("- Manage concurrent patient admissions using processes and threads.", style='List Bullet')
    doc.add_paragraph("- Implement memory management strategies for bed allocation.", style='List Bullet')
    doc.add_paragraph("- Visualize the hospital ward in real-time.", style='List Bullet')

    # 4. System Design
    doc.add_heading('2. System Design', level=1)
    doc.add_paragraph("The system relies on POSIX standards and runs on Linux (WSL). It is split into two main layers:")
    doc.add_paragraph("1. Admission Layer: Bash scripts handle initial intake, compute triage priority (1-5), and pipe records to the C binary.", style='List Number')
    doc.add_paragraph("2. Hospital Layer: A multithreaded C application runs the hospital simulation. It includes a receptionist thread, a scheduler thread, and a nurse thread pool.", style='List Number')

    # 5. Phase 1 Docs
    doc.add_heading('3. Phase 1: Shell Scripts', level=1)
    doc.add_paragraph("We built three shell scripts to handle setup and intake.")
    doc.add_paragraph("triage.sh: Takes name, age, and severity (1-10) as arguments. It calculates a triage priority from 1 (Critical) to 5 (Minimal) and outputs a formatted patient record.")
    doc.add_paragraph("start_hospital.sh: Prepares named FIFOs and launches the admissions manager in the background.")
    doc.add_paragraph("stress_test.sh: Automates 25 patient arrivals to test queue limits and thread safety under heavy load.")

    # 6. Phase 2 Docs
    doc.add_heading('4. Phase 2: Process Management & IPC', level=1)
    doc.add_paragraph("The admissions manager uses fork() and execv() to spawn a patient_simulator process for every admitted patient. To prevent zombie processes, the admissions manager catches SIGCHLD and reaps children using waitpid(WNOHANG).")
    doc.add_paragraph("Processes communicate through named pipes (FIFOs). The triage script pipes data to admissions. When treatment finishes, patient_simulator writes its ID to a discharge FIFO. The ward bitmap is stored in shared memory so all processes can see the bed states.")
    
    # 7. Phase 3 Docs
    doc.add_heading('5. Phase 3: Threads & Synchronization', level=1)
    doc.add_paragraph("We implemented a priority queue protected by a mutex (queue_mutex) and condition variable (queue_cond).")
    doc.add_paragraph("- Receptionist Thread: Reads the triage FIFO and adds patients to the queue.")
    doc.add_paragraph("- Scheduler Thread: Dequeues the highest-priority patient, finds a bed, and spawns the patient process.")
    doc.add_paragraph("- Nurse Threads (ICU, Isolation, General): Wait for discharge signals. When a patient leaves, the nurse frees the bed and broadcasts a bed_freed signal so the scheduler can wake up.")

    # 8. Phase 4 Docs
    doc.add_heading('6. Phase 4: Memory Management', level=1)
    doc.add_paragraph("The 20 hospital beds map to 32 discrete 'care units' stored in a contiguous array. We implemented Best-Fit, First-Fit, and Worst-Fit allocation strategies.")
    doc.add_paragraph("We track the free list and calculate external fragmentation after every admission. A real bug we encountered was attempting to coalesce adjacent free beds. Since physical beds have fixed identities, coalescing erased beds from the ward. We fixed this by making beds fixed-size partitions. We also log memory stats to memory_log.txt and use mmap() to persist patient records.")

    # 9. Testing
    doc.add_heading('7. Testing', level=1)
    doc.add_paragraph("We used stress_test.sh to verify thread safety. The test sends 20 patients normally, followed by a burst of 5 high-priority patients. The UI correctly handles queue depths up to 18, and beds are accurately allocated and freed without crashing or race conditions.")

    # 10. Valgrind Output
    doc.add_heading('8. Valgrind Output', level=1)
    doc.add_paragraph("The project compiles with zero warnings under -Wall -Wextra. Valgrind tests show no memory leaks upon clean exit. All shared memory and named semaphores are successfully unlinked by stop_hospital.sh.")

    # 11. Challenges & Lessons Learned
    doc.add_heading('9. Challenges & Lessons Learned', level=1)
    doc.add_paragraph("1. Startup Deadlocks: Opening an empty FIFO with O_RDONLY blocks the thread. We solved this by having the admissions thread open the FIFO with O_RDWR so it never blocks.")
    doc.add_paragraph("2. Coalescing Bug: Applying heap allocator logic (coalescing) to a physical resource (beds) caused slots to disappear. We learned that domain logic always dictates how OS concepts should be applied.")
    doc.add_paragraph("3. WSL Line Endings: Bash scripts written in Windows failed to run in WSL due to CRLF line endings. We added a .gitattributes file to force LF endings globally.")

    # 12. Individual Contribution
    doc.add_heading('10. Individual Contribution', level=1)
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Roll Number'
    hdr_cells[2].text = 'Key Contributions'
    
    row_cells = table.rows[1].cells
    row_cells[0].text = 'Abdul Ahad Zawiar'
    row_cells[1].text = '24F-0727'
    row_cells[2].text = 'Bed Allocator, Terminal UI, Memory Logging, Threads'
    
    row_cells = table.rows[2].cells
    row_cells[0].text = 'Member 2'
    row_cells[1].text = '24F-0514'
    row_cells[2].text = 'Shell Scripts, IPC Setup, Paging, Documentation'

    # Placeholder for the Video
    doc.add_heading('Demo Video', level=1)
    doc.add_paragraph('[ PASTE YOUR DEMO VIDEO LINK OR EMBED HERE ]')

    doc.save('report/report.docx')
    print("report.docx generated successfully in the report folder.")

if __name__ == '__main__':
    main()
